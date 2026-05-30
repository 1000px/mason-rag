from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
)

SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".xlsx": "excel",
    ".xls": "excel",
    ".docx": "docx",
}


class DocumentLoader:
    def __init__(self, file_path: str | Path):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {self.file_path}")
        ext = self.file_path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}，支持: {list(SUPPORTED_EXTENSIONS.keys())}")
        self.doc_type = SUPPORTED_EXTENSIONS[ext]

    def load(self) -> List[Document]:
        if self.doc_type == "pdf":
            return self._load_pdf()
        elif self.doc_type == "text":
            return self._load_text()
        elif self.doc_type == "markdown":
            return self._load_markdown()
        elif self.doc_type == "excel":
            return self._load_excel()
        elif self.doc_type == "docx":
            return self._load_docx()
        return []

    def _load_pdf(self) -> List[Document]:
        loader = PyPDFLoader(str(self.file_path))
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = str(self.file_path)
            doc.metadata["file_type"] = "pdf"
            doc.metadata["file_name"] = self.file_path.name
        return docs

    def _load_text(self) -> List[Document]:
        loader = TextLoader(str(self.file_path), encoding="utf-8")
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = str(self.file_path)
            doc.metadata["file_type"] = "text"
            doc.metadata["file_name"] = self.file_path.name
        return docs

    def _load_markdown(self) -> List[Document]:
        loader = TextLoader(str(self.file_path), encoding="utf-8")
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = str(self.file_path)
            doc.metadata["file_type"] = "markdown"
            doc.metadata["file_name"] = self.file_path.name
        return docs

    def _load_excel(self) -> List[Document]:
        ext = self.file_path.suffix.lower()

        if ext == ".xlsx":
            return self._load_xlsx()
        elif ext == ".xls":
            return self._load_xls()
        else:
            raise ValueError(f"不支持的Excel格式: {ext}")

    def _load_xlsx(self) -> List[Document]:
        import openpyxl

        wb = openpyxl.load_workbook(str(self.file_path), data_only=True)
        docs: List[Document] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([
                    str(cell) if cell is not None else ""
                    for cell in row
                ]).strip(" |")
                if row_text:
                    rows.append(row_text)
            if not rows:
                continue
            text = f"# Sheet: {sheet_name}\n\n" + "\n".join(rows)
            docs.append(Document(page_content=text, metadata={
                "source": str(self.file_path),
                "file_type": "excel",
                "file_name": self.file_path.name,
                "sheet_name": sheet_name,
            }))
        return docs

    def _load_xls(self) -> List[Document]:
        import xlrd

        wb = xlrd.open_workbook(str(self.file_path))
        docs: List[Document] = []
        for sheet in wb.sheets():
            rows = []
            for row_idx in range(sheet.nrows):
                row_text = " | ".join([
                    str(sheet.cell_value(row_idx, col_idx))
                    for col_idx in range(sheet.ncols)
                ]).strip(" |")
                if row_text:
                    rows.append(row_text)
            if not rows:
                continue
            text = f"# Sheet: {sheet.name}\n\n" + "\n".join(rows)
            docs.append(Document(page_content=text, metadata={
                "source": str(self.file_path),
                "file_type": "excel",
                "file_name": self.file_path.name,
                "sheet_name": sheet.name,
            }))
        return docs

    def _load_docx(self) -> List[Document]:
        import docx

        doc = docx.Document(str(self.file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([
                    cell.text for cell in row.cells if cell.text.strip()
                ])
                if row_text.strip():
                    tables_text.append(row_text)
        parts = paragraphs
        if tables_text:
            parts.append("\n--- 表格内容 ---\n")
            parts.extend(tables_text)
        text = "\n".join(parts)
        return [Document(page_content=text, metadata={
            "source": str(self.file_path),
            "file_type": "docx",
            "file_name": self.file_path.name,
        })]


def load_document(file_path: str | Path) -> List[Document]:
    loader = DocumentLoader(file_path)
    return loader.load()